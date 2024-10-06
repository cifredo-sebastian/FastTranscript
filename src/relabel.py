import tkinter as tk
from tkinter import filedialog, messagebox, ttk
import docx
import os
import sys

def relabel(file_path,filetype,status_label, file_destination):
    if file_destination:
        speaker_count, speaker_ids = count_speakers(file_path,filetype)
        if speaker_count == 0:
            messagebox.showinfo("Labels not found", f"No speaker labels found in {file_path}.")
        else:
            open_relabler(file_path,filetype,status_label,speaker_ids, speaker_count, file_destination)

def open_relabler(file_path, filetype, status_label, speaker_ids, speaker_count, file_destination):
    relabel_window = tk.Toplevel()
    relabel_window.title("Relabel")
    if hasattr(sys, '_MEIPASS'):
        icon_path = os.path.join(sys._MEIPASS, 'public', 'fasttranscript.ico')
    else:
        icon_path = 'public/fasttranscript.ico'
    relabel_window.iconbitmap(icon_path)

    relabel_window.geometry("300x400")

    container = ttk.Frame(relabel_window)
    container.pack(fill='both', expand=True)

    canvas = tk.Canvas(container, width=280)
    canvas.pack(side='left', fill='both', expand=True)

    scrollbar = ttk.Scrollbar(container, orient='vertical', command=canvas.yview)
    scrollbar.pack(side='right', fill='y')

    canvas.configure(yscrollcommand=scrollbar.set)
    
    canvas.bind('<Configure>', lambda e: canvas.configure(scrollregion=canvas.bbox("all")))

    scrollable_frame = ttk.Frame(canvas)
    
    scrollable_frame.bind(
        "<Configure>",
        lambda e: canvas.itemconfig("window", width=canvas.winfo_width())
    )
    
    canvas.create_window((0, 0), window=scrollable_frame, anchor='nw', tags="window")

    entries = {}

    # Add speaker labels and entry fields
    for speaker in speaker_ids:
        label = tk.Label(scrollable_frame, text=f'{speaker}: ')
        label.pack(anchor='w', padx=10, pady=5)
        entry = tk.Entry(scrollable_frame)
        entry.pack(fill='x', padx=10, pady=5)

        entries[speaker] = entry

    def save_data():
        speaker_data = {speaker: entries[speaker].get() for speaker in speaker_ids}
        speaker_relabel(file_path, filetype, speaker_data, file_destination)
        messagebox.showinfo("Relabel Saved", f"File relabeled and saved in {file_destination}.")
        relabel_window.destroy()

    tk.Button(scrollable_frame, text="Save", command=save_data).pack(padx=10, pady=10)



def speaker_relabel(file_path, filetype, speaker_data,file_destination):
    if filetype == 'docx':
        doc = docx.Document(file_path)
        for para in doc.paragraphs:
            lines = para.text.splitlines()
            updated_lines = []
            for line in lines:
                for old_speaker, new_speaker in speaker_data.items():
                    if old_speaker in line:
                        line = line.replace(old_speaker, new_speaker)
                updated_lines.append(line)

            para.text = "\n".join(updated_lines)
        doc.save(file_destination)

    else:
        with open(file_path, 'r') as file:
            file_content = file.read()

        for old_speaker, new_speaker in speaker_data.items():
            file_content = file_content.replace(old_speaker, new_speaker)

        with open(file_destination, 'w') as file:
            file.write(file_content)

    print(f"File saved to {file_destination}")


def count_speakers(file_path,filetype):
    speakers = set()

    if filetype == 'docx':
        doc = docx.Document(file_path)
        for para in doc.paragraphs:
            lines = para.text.splitlines()
            for line in lines:
                if 'Speaker' in line:
                    start = line.find('Speaker ') + len('Speaker ')
                    end = line.find(':', start)
                    speaker_id = "Speaker " + line[start:end].strip()
                    speakers.add(speaker_id)

    else:
        with open(file_path, 'r') as file:
            for line in file:
                if 'Speaker' in line:
                    start = line.find('Speaker ') + len('Speaker ')
                    end = line.find(':', start)
                    speaker_id = "Speaker " + line[start:end].strip()
                    speakers.add(speaker_id)

    return len(speakers), sorted(speakers)
